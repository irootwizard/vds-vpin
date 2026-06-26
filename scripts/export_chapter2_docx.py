#!/usr/bin/env python3
"""Export docs/vPIN-作品设计与实现.md to Word using the official reference template."""

from __future__ import annotations

import re
import shutil
import subprocess
import sys
import os
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
SRC_MD = DOCS / "vPIN-作品设计与实现.md"
EXPORT_MD = DOCS / "vPIN-作品设计与实现-export.md"
FIGURES = DOCS / "figures"
TEMPLATE = DOCS / "【作品报告】面向CNN隐私推理的同态加密与零知识可验证计算系统：CipherVision.docx"
OUTPUT = DOCS / "第二章-作品设计与实现.docx"
BUILD_OUTPUT = DOCS / "_chapter2_build.docx"

MERMAID_MAP = [
    "fig-2-1-系统逻辑框架",
    "fig-2-2-系统总体架构",
    "fig-2-3-系统网络拓扑",
    "fig-2-4-软件模块依赖",
    "fig-2-6-协议时序",
    "fig-2-7-会话状态机",
    "fig-2-8-绑定闭合链",
]

FIGURE_TITLES = {
    "fig-2-1-系统逻辑框架": "图 2-1 系统逻辑框架（双平面）",
    "fig-2-2-系统总体架构": "图 2-2 系统总体架构（三层）",
    "fig-2-3-系统网络拓扑": "图 2-3 系统网络拓扑",
    "fig-2-4-软件模块依赖": "图 2-4 软件模块依赖关系",
    "fig-2-5-工作流程泳道": "图 2-5 系统工作流程图（参与方泳道）",
    "fig-2-6-协议时序": "图 2-6 端到端协议时序",
    "fig-2-7-会话状态机": "图 2-7 推理会话状态机",
    "fig-2-8-绑定闭合链": "图 2-8 密码学绑定闭合链",
    "fig-2-9-Merkle稀疏打开": "图 2-9 大模型 Merkle 稀疏打开",
    "fig-2-10-客户端界面": "图 2-10 客户端界面与协议阶段",
}


def run(cmd: list[str], cwd: Path | None = None) -> None:
    print("+", " ".join(cmd))
    subprocess.run(cmd, cwd=cwd, check=True)


def extract_mermaid_blocks(text: str) -> list[str]:
    pattern = re.compile(r"```mermaid\s*\n(.*?)```", re.DOTALL)
    return [m.group(1).strip() for m in pattern.finditer(text)]


def preprocess_markdown(text: str, *, include_images: bool = False) -> str:
    # Remove editor-only figure index table at chapter start
    text = re.sub(
        r"\*\*本章插图一览.*?\n\n---\n\n",
        "",
        text,
        count=1,
        flags=re.DOTALL,
    )

    # Remove illustration instruction blockquotes
    text = re.sub(
        r"> \*\*【插图说明】\*\*.*?(?=\n\n(?:!\[|<!--|```|\*\*图))",
        "",
        text,
        flags=re.DOTALL,
    )

    # Remove HTML comments and mermaid source blocks (images remain above)
    text = re.sub(r"<!--.*?-->\s*\n", "", text, flags=re.DOTALL)
    text = re.sub(r"```mermaid\s*\n.*?```\s*\n", "", text, flags=re.DOTALL)

    # Remove embedded images unless explicitly requested
    if not include_images:
        text = re.sub(r"!\[[^\]]*\]\([^)]+\)\s*\n", "", text)

    # Pandoc-friendly math: drop \tag{n}
    text = re.sub(r"\\tag\{\d+\}", "", text)

    # Footer version line
    text = re.sub(r"\n---\n\n\*文档版本：.*\n?$", "\n", text)

    # Horizontal rules cause extra breaks in Word
    text = re.sub(r"\n---\n", "\n", text)

    # Pandoc table captions (caption above table)
    text = re.sub(
        r"\*\*(表 2-\d+[^*]+)\*\*\s*\n+(\|)",
        r"Table: \1\n\n\2",
        text,
    )

    # Figure captions: plain centered line (no bold markdown)
    text = re.sub(
        r"\*\*(图 2-\d+[^*]+)\*\*",
        r"\1",
        text,
    )

    # Collapse excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip() + "\n"


def export_mermaid_png(mmd_path: Path, png_path: Path) -> bool:
    npx = shutil.which("npx")
    if not npx:
        return False
    try:
        subprocess.run(
            [
                npx,
                "-y",
                "@mermaid-js/mermaid-cli",
                "-i",
                str(mmd_path),
                "-o",
                str(png_path),
                "-b",
                "white",
                "-w",
                "1400",
            ],
            cwd=FIGURES,
            check=True,
            timeout=120,
            capture_output=True,
            text=True,
        )
        return png_path.exists()
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
        print(f"! mermaid export failed for {mmd_path.name}: {exc}", file=sys.stderr)
        return False


def make_placeholder_png(png_path: Path, title: str, subtitle: str = "【插图占位，可替换】") -> None:
    from PIL import Image, ImageDraw, ImageFont

    img = Image.new("RGB", (1200, 700), color=(248, 248, 248))
    draw = ImageDraw.Draw(img)
    draw.rectangle([(40, 40), (1160, 660)], outline=(160, 160, 160), width=3)
    try:
        font = ImageFont.truetype("msyh.ttc", 36)
        font_small = ImageFont.truetype("msyh.ttc", 26)
    except OSError:
        font = ImageFont.load_default()
        font_small = font

    for i, line in enumerate([title, subtitle]):
        f = font if i == 0 else font_small
        bbox = draw.textbbox((0, 0), line, font=f)
        w = bbox[2] - bbox[0]
        draw.text(((1200 - w) / 2, 300 + i * 48), line, fill=(90, 90, 90), font=f)
    img.save(png_path)


def postprocess_docx(docx_path: Path) -> None:
    """Fix heading levels, captions, spacing, and table borders for Word output."""
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    doc = Document(str(docx_path))

    # Ensure Heading 3 exists (reference template only has Heading 1/2)
    if "Heading 3" not in [s.name for s in doc.styles]:
        h3 = doc.styles.add_style("Heading 3", WD_STYLE_TYPE.PARAGRAPH)
        h2 = doc.styles["Heading 2"]
        h3.base_style = h2
        h3.font.name = h2.font.name
        h3.font.size = Pt(14)
        h3.font.bold = True

    h3_re = re.compile(r"^\d+\.\d+\.\d+\s+\S")
    h2_re = re.compile(r"^\d+\.\d+\s+\S")
    fig_re = re.compile(r"^图 2-\d+")
    tbl_re = re.compile(r"^表 2-\d+")

    prev_empty = False
    to_remove = []
    for para in doc.paragraphs:
        text = para.text.strip()

        if not text:
            if prev_empty:
                to_remove.append(para)
            prev_empty = True
            continue
        prev_empty = False

        if h3_re.match(text):
            para.style = doc.styles["Heading 3"]
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.keep_with_next = True
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = 1.5
        elif h2_re.match(text) and not h3_re.match(text):
            if para.style.name not in ("Heading 1", "Heading 2"):
                para.style = doc.styles["Heading 2"]
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.keep_with_next = True
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = 1.5
        elif fig_re.match(text) or tbl_re.match(text):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10.5)

    for para in to_remove:
        p = para._element
        p.getparent().remove(p)

    def set_cell_border(cell, **kwargs):
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn("w:tcBorders"))
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        for edge, border_attrs in kwargs.items():
            tag = f"w:{edge}"
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key, value in border_attrs.items():
                element.set(qn(f"w:{key}"), str(value))

    border = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(
                    cell,
                    top=border,
                    bottom=border,
                    left=border,
                    right=border,
                )
        try:
            table.style = "Table Grid"
        except KeyError:
            pass

    for para in doc.paragraphs:
        if para.style.name == "Heading 1":
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = 1.5

    doc.save(str(docx_path))


def main() -> int:
    if not SRC_MD.exists():
        print(f"Missing source: {SRC_MD}", file=sys.stderr)
        return 1
    if not TEMPLATE.exists():
        print(f"Missing template: {TEMPLATE}", file=sys.stderr)
        return 1

    FIGURES.mkdir(parents=True, exist_ok=True)
    source = SRC_MD.read_text(encoding="utf-8")
    include_images = os.environ.get("EXPORT_INCLUDE_IMAGES", "0") == "1"

    if include_images:
        blocks = extract_mermaid_blocks(source)
        if len(blocks) != len(MERMAID_MAP):
            print(
                f"Warning: expected {len(MERMAID_MAP)} mermaid blocks, found {len(blocks)}",
                file=sys.stderr,
            )
        skip_mermaid = os.environ.get("EXPORT_SKIP_MERMAID", "1") == "1"
        for name, block in zip(MERMAID_MAP, blocks):
            mmd = FIGURES / f"{name}.mmd"
            png = FIGURES / f"{name}.png"
            mmd.write_text(block + "\n", encoding="utf-8")
            if png.exists():
                continue
            if skip_mermaid or not export_mermaid_png(mmd, png):
                make_placeholder_png(
                    png,
                    FIGURE_TITLES.get(name, name),
                    "【运行 EXPORT_SKIP_MERMAID=0 可从 Mermaid 渲染】"
                    if skip_mermaid
                    else "【Mermaid 导出失败，占位图】",
                )
        for name, title in FIGURE_TITLES.items():
            png = FIGURES / f"{name}.png"
            if not png.exists():
                sub = "【待手绘】" if name in {
                    "fig-2-5-工作流程泳道",
                    "fig-2-9-Merkle稀疏打开",
                    "fig-2-10-客户端界面",
                } else "【Mermaid 导出失败，占位图】"
                make_placeholder_png(png, title, sub)
    else:
        print("Skip figures (EXPORT_INCLUDE_IMAGES=0): Word 仅保留图题，不嵌入图片")

    export_text = preprocess_markdown(source, include_images=include_images)
    EXPORT_MD.write_text(export_text, encoding="utf-8")

    pandoc = shutil.which("pandoc")
    if not pandoc:
        print("pandoc not found", file=sys.stderr)
        return 1

    run(
        [
            pandoc,
            str(EXPORT_MD),
            "-o",
            str(BUILD_OUTPUT),
            "--from",
            "markdown+tex_math_dollars",
            "--standalone",
            f"--resource-path={DOCS}",
            f"--reference-doc={TEMPLATE}",
            "--metadata",
            "lang=zh-CN",
        ],
        cwd=DOCS,
    )

    postprocess_docx(BUILD_OUTPUT)

    try:
        if OUTPUT.exists():
            OUTPUT.unlink()
        BUILD_OUTPUT.replace(OUTPUT)
        final_path = OUTPUT
    except PermissionError:
        alt = DOCS / "第二章-作品设计与实现-新版.docx"
        if alt.exists():
            alt.unlink()
        BUILD_OUTPUT.replace(alt)
        final_path = alt
        print(f"! 原文件被占用，已写入: {alt}", file=sys.stderr)

    print(f"\nDone: {final_path}")
    if include_images:
        print(f"Figures: {FIGURES}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
